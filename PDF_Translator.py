import importlib
import subprocess
import sys
def import_or_install(package, import_name=None):
    """
    Eğer kütüphane yüklü değilse pip ile yükler, sonra import eder.
    package: pip install için kullanılacak isim
    import_name: import için kullanılacak isim (farklıysa)
    """
    try:
        return importlib.import_module(import_name or package)
    except ImportError:
        print(f"[INFO] {package} bulunamadı, yükleniyor...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
        return importlib.import_module(import_name or package)
docx = import_or_install("python-docx", "docx")
PyPDF2 = import_or_install("PyPDF2")
requests = import_or_install("requests")
import os
import threading
import math
import time
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from pathlib import Path
import requests
from docx import Document
from PyPDF2 import PdfReader


# ---------- Ayarlar ----------
# DeepL endpoint (global). Eğer PRO kullanıyorsan endpoint farklı olabilir.
DEEPL_API_URL = "https://api-free.deepl.com/v2/translate"  # free endpoint
CHUNK_CHAR_SIZE = 4000  # (güvenli bir varsayılan)


LANGUAGES = [
    "AUTO", "BG", "CS", "DA", "DE", "EL", "EN", "ES", "ET", "FI",
    "FR", "HU", "ID", "IT", "JA", "LT", "LV", "NB", "NL", "PL",
    "PT", "RO", "RU", "SK", "SL", "SV", "TR", "UK", "ZH"
]

# ---------- Yardımcı Fonksiyonlar ----------
def human_readable_size(bytesize: int) -> str:
    if bytesize < 1024:
        return f"{bytesize} B"
    for unit in ("KB","MB","GB","TB"):
        bytesize /= 1024.0
        if bytesize < 1024.0:
            return f"{bytesize:.2f} {unit}"
    return f"{bytesize:.2f} PB"

def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    texts = []
    for p in reader.pages:
        txt = p.extract_text()
        if txt:
            texts.append(txt)
    return "\n\n".join(texts)

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n\n".join(paragraphs)

def get_file_info(path: str) -> dict:
    p = Path(path)
    size = p.stat().st_size
    ext = p.suffix.lower()
    page_count = 0
    word_count = 0
    try:
        if ext == ".pdf":
            reader = PdfReader(path)
            page_count = len(reader.pages)
            text = extract_text_from_pdf(path)
            word_count = len(text.split())
        elif ext == ".docx":
            doc = Document(path)
            # docx page count isn't trivial; use paragraphs length as naive proxy
            text = extract_text_from_docx(path)
            word_count = len(text.split())
            page_count = "-"  # gerçek sayfa sayısını docx ile kolayca almak zordur
        else:
            text = ""
    except Exception as e:
        text = ""
    return {
        "name": p.name,
        "size": size,
        "size_human": human_readable_size(size),
        "ext": ext,
        "page_count": page_count,
        "word_count": word_count,
    }

def chunk_text_preserve_paragraphs(text: str, chunk_size: int):
    """
    Paragrafları koruyarak yaklaşık chunk_size karaktere böler.
    Döndürür: list of strings
    """
    paragraphs = [p for p in text.split("\n") if p.strip() != ""]
    chunks = []
    current = ""
    for p in paragraphs:
        if len(current) + len(p) + 1 <= chunk_size:
            current += (("\n" if current else "") + p)
        else:
            if current:
                chunks.append(current)
            # Eğer tek paragraf chunk_size'dan uzun ise onu da parçala
            if len(p) > chunk_size:
                # kaba parça
                for i in range(0, len(p), chunk_size):
                    chunks.append(p[i:i+chunk_size])
                current = ""
            else:
                current = p
    if current:
        chunks.append(current)
    return chunks

def call_deepl_translate(text: str, source_lang: str, target_lang: str, api_key: str) -> str:
    """
    Basit DeepL çağrısı. source_lang "AUTO" ise source paramı gönderilmez.
    """
    if not api_key:
        raise ValueError("DeepL API key gerekli.")
    data = {
        "auth_key": api_key,
        "text": text,
        "target_lang": target_lang
    }
    if source_lang and source_lang.upper() != "AUTO":
        data["source_lang"] = source_lang
    # DeepL supports multiple text fields but we'll send one chunk per request
    resp = requests.post(DEEPL_API_URL, data=data, timeout=60)
    if resp.status_code != 200:
        raise RuntimeError(f"DeepL API hatası: {resp.status_code} - {resp.text}")
    j = resp.json()
    # yanıt structure: {"translations":[{"detected_source_language":"EN","text":"..."}]}
    translated = " ".join(t["text"] for t in j.get("translations", []))
    return translated

def save_text_as_docx(text: str, out_path: str):
    doc = Document()
    # basit: paragraf kırılmalarını koru
    for para in text.split("\n"):
        doc.add_paragraph(para)
    doc.save(out_path)

# ---------- GUI Uygulaması ----------
class TranslatorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("DeepL Çeviri - Tkinter")
        self.root.geometry("760x520")
        self.root.minsize(700, 480)

        # Stil
        self.style = ttk.Style(self.root)
        try:
            self.style.theme_use("clam")
        except:
            pass

        # Üst kısım: API key + Dosya seçimi
        top_frame = ttk.Frame(root, padding=10)
        top_frame.pack(fill="x")

        ttk.Label(top_frame, text="DeepL API Key:").grid(row=0, column=0, sticky="w")
        self.api_entry = ttk.Entry(top_frame, width=60, show="*")
        self.api_entry.grid(row=0, column=1, sticky="w", padx=(6,6))
        self.show_key_var = tk.BooleanVar(value=False)
        self.show_key_cb = ttk.Checkbutton(top_frame, text="Göster", variable=self.show_key_var, command=self.toggle_api_show)
        self.show_key_cb.grid(row=0, column=2, sticky="w")

        ttk.Button(top_frame, text="Dosya Seç (.pdf veya .docx)", command=self.select_file).grid(row=1, column=0, columnspan=3, pady=(10,0), sticky="w")

        # Dil seçimleri
        langs_frame = ttk.Frame(root, padding=10)
        langs_frame.pack(fill="x")

        ttk.Label(langs_frame, text="Kaynak Dil:").grid(row=0, column=0, sticky="w")
        self.src_combo = ttk.Combobox(langs_frame, values=LANGUAGES, width=12)
        self.src_combo.set("AUTO")
        self.src_combo.grid(row=0, column=1, padx=(6,20), sticky="w")

        ttk.Label(langs_frame, text="Hedef Dil:").grid(row=0, column=2, sticky="w")
        self.tgt_combo = ttk.Combobox(langs_frame, values=LANGUAGES, width=12)
        self.tgt_combo.set("TR")
        self.tgt_combo.grid(row=0, column=3, padx=(6,20), sticky="w")

        # Orta kısım: Dosya bilgileri
        info_frame = ttk.LabelFrame(root, text="Dosya Bilgileri", padding=10)
        info_frame.pack(fill="x", padx=10, pady=10)

        self.file_path_var = tk.StringVar(value="Henüz dosya seçilmedi.")
        ttk.Label(info_frame, textvariable=self.file_path_var).pack(anchor="w")

        self.file_info_var = tk.StringVar(value="")
        ttk.Label(info_frame, textvariable=self.file_info_var).pack(anchor="w", pady=(6,0))

        # Progress ve durumu
        progress_frame = ttk.Frame(root, padding=10)
        progress_frame.pack(fill="x")

        self.progress = ttk.Progressbar(progress_frame, length=520, mode="determinate")
        self.progress.pack(side="left", padx=(0,10))
        self.progress["value"] = 0

        self.status_var = tk.StringVar(value="Bekleniyor...")
        ttk.Label(progress_frame, textvariable=self.status_var).pack(side="left")

        # Butonlar
        btn_frame = ttk.Frame(root, padding=10)
        btn_frame.pack(fill="x")

        self.translate_btn = ttk.Button(btn_frame, text="Çevir", command=self.on_translate_clicked)
        self.translate_btn.pack(side="left")

        self.open_btn = ttk.Button(btn_frame, text="Çevrilen Dosyayı Aç", state="disabled", command=self.open_translated)
        self.open_btn.pack(side="left", padx=(10,0))

        self.clear_btn = ttk.Button(btn_frame, text="Temizle", command=self.clear_all)
        self.clear_btn.pack(side="right")

        # Metin gösterimi (scrollable)
        out_frame = ttk.LabelFrame(root, text="Önizleme (Çevrilen metin kaydedildikten sonra yüklenir)", padding=6)
        out_frame.pack(fill="both", expand=True, padx=10, pady=10)

        self.text_widget = tk.Text(out_frame, wrap="word")
        self.text_widget.pack(fill="both", expand=True, side="left")
        scrollbar = ttk.Scrollbar(out_frame, command=self.text_widget.yview)
        scrollbar.pack(side="right", fill="y")
        self.text_widget.config(yscrollcommand=scrollbar.set)

        # internal state
        self.selected_file = None
        self.translated_path = None
        self._worker_thread = None
        self._stop_requested = False

    def toggle_api_show(self):
        if self.show_key_var.get():
            self.api_entry.config(show="")
        else:
            self.api_entry.config(show="*")

    def select_file(self):
        fp = filedialog.askopenfilename(
            title="PDF veya DOCX seç",
            filetypes=[("PDF Files", "*.pdf"), ("Word Documents", "*.docx")]
        )
        if not fp:
            return
        self.selected_file = fp
        finfo = get_file_info(fp)
        self.file_path_var.set(f"Seçili: {finfo['name']}")
        info_text = f"Uzantı: {finfo['ext']} | Boyut: {finfo['size_human']} | Sayfa: {finfo['page_count']} | Kelime (tahmini): {finfo['word_count']}"
        self.file_info_var.set(info_text)
        self.status_var.set("Dosya seçildi. Çevirmek için 'Çevir' butonuna bas.")
        self.text_widget.delete("1.0", tk.END)
        self.open_btn.config(state="disabled")
        self.translated_path = None

    def on_translate_clicked(self):
        if not self.selected_file:
            messagebox.showerror("Hata", "Lütfen önce bir dosya seç.")
            return
        api_key = self.api_entry.get().strip()
        if not api_key:
            if not messagebox.askyesno("API Key yok", "DeepL API anahtarı girilmedi. Devam etmek istiyor musun? (bu durumda istek başarısız olur)"):
                return
        src = self.src_combo.get().strip()
        tgt = self.tgt_combo.get().strip()
        if not tgt:
            messagebox.showerror("Hata", "Lütfen hedef dil seç.")
            return
        # disable UI
        self.translate_btn.config(state="disabled")
        self._stop_requested = False
        # start worker thread
        self._worker_thread = threading.Thread(target=self._translate_worker, args=(self.selected_file, src, tgt, api_key), daemon=True)
        self._worker_thread.start()

    def _translate_worker(self, file_path, src, tgt, api_key):
        try:
            self._set_status("Dosya okunuyor...")
            ext = Path(file_path).suffix.lower()
            if ext == ".pdf":
                text = extract_text_from_pdf(file_path)
            elif ext == ".docx":
                text = extract_text_from_docx(file_path)
            else:
                raise ValueError("Desteklenmeyen dosya türü.")

            if not text.strip():
                raise ValueError("Dosyadan metin alınamadı veya içerik boş.")

            # chunk the text
            chunks = chunk_text_preserve_paragraphs(text, CHUNK_CHAR_SIZE)
            total = len(chunks)
            self._set_progress(0, total)
            translated_parts = []
            for idx, ch in enumerate(chunks, start=1):
                if self._stop_requested:
                    self._set_status("İptal edildi.")
                    return
                self._set_status(f"Çeviriliyor... ({idx}/{total})")
                try:
                    translated = call_deepl_translate(ch, src, tgt, api_key)
                except Exception as e:
                    # hata durumunda durdurup kullanıcıya bildir
                    raise RuntimeError(f"Çeviri sırasında hata: {e}")
                translated_parts.append(translated)
                self._set_progress(idx, total)
                # kısa bekleme (DeepL rate limit'ine karşı tedbir)
                time.sleep(0.1)

            full_translated = "\n\n".join(translated_parts)
            # kaydet
            out_path = Path(file_path).with_name(Path(file_path).stem + "_translated.docx")
            save_text_as_docx(full_translated, str(out_path))
            self.translated_path = str(out_path)
            # GUI güncelle
            self._set_status("Çeviri tamamlandı. Kayıt yapıldı.")
            self._set_progress(1,1)  # tamam
            # yükle önizlemeye (ilk 2000 karakter)
            preview = full_translated[:100000]  # büyük dosyalarda sınır
            self._set_text_widget(preview)
            # enable open button
            self.root.after(0, lambda: self.open_btn.config(state="normal"))
            messagebox.showinfo("Bitti", f"Çeviri dosyası kaydedildi:\n{out_path}")
        except Exception as e:
            messagebox.showerror("Hata", str(e))
            self._set_status("Hata oluştu.")
        finally:
            self.translate_btn.config(state="normal")

    def _set_progress(self, value, total):
        # progress value between 0-100
        try:
            pct = int((value/total) * 100) if total and total>0 else 0
        except:
            pct = 0
        self.root.after(0, lambda: self.progress.config(value=pct))

    def _set_status(self, text):
        self.root.after(0, lambda: self.status_var.set(text))

    def _set_text_widget(self, content):
        def inner():
            self.text_widget.delete("1.0", tk.END)
            self.text_widget.insert("1.0", content)
        self.root.after(0, inner)

    def open_translated(self):
        if not self.translated_path or not Path(self.translated_path).exists():
            messagebox.showerror("Hata", "Çevrilen dosya bulunamadı.")
            return
        p = Path(self.translated_path)
        try:
            if os.name == "nt":
                os.startfile(str(p))
            elif os.name == "posix":
                # mac veya linux
                if "darwin" in os.sys.platform:
                    os.system(f"open '{p}'")
                else:
                    os.system(f"xdg-open '{p}'")
            else:
                messagebox.showinfo("Bilgi", f"Dosya konumu: {p}")
        except Exception as e:
            messagebox.showerror("Hata", f"Dosya açılamadı: {e}")

    def clear_all(self):
        if messagebox.askyesno("Temizle", "Formu temizlemek istiyor musun?"):
            self.api_entry.delete(0, tk.END)
            self.file_path_var.set("Henüz dosya seçilmedi.")
            self.file_info_var.set("")
            self.status_var.set("Bekleniyor...")
            self.progress.config(value=0)
            self.text_widget.delete("1.0", tk.END)
            self.selected_file = None
            self.translated_path = None
            self.open_btn.config(state="disabled")


if __name__ == "__main__":


    root = tk.Tk()
    app = TranslatorApp(root)
    root.mainloop()
